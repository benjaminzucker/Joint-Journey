#!/usr/bin/env python3
"""
Regenerate: Joint Journey - Intended Use and Safety Summary.docx

Produces a valid, Word-openable .docx of the Intended Use Statement and
Safety Summary for Joint Journey. Run:

    python3 generate_intended_use_docx.py

Requires: python-docx  (pip3 install python-docx)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Joint Journey - Intended Use and Safety Summary.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def title(doc, text, size, color=NAVY, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def meta(doc, pairs):
    p = doc.add_paragraph()
    for i, (label, value) in enumerate(pairs):
        if i:
            p.add_run("    ")
        p.add_run(label + " ").bold = True
        p.add_run(value)
    p.paragraph_format.space_after = Pt(10)
    return p


def para(doc, text, after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(after)
    return p


def statement(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(8)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.size = Pt(11)


def bullets_status(doc, items):
    """items = list of (text, status) — status shown bold at the front."""
    for text, status in items:
        p = doc.add_paragraph(style="List Bullet")
        s = p.add_run(f"[{status}] ")
        s.bold = True
        s.font.color.rgb = NAVY
        s.font.size = Pt(11)
        r = p.add_run(text)
        r.font.size = Pt(11)


def build():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- Header ----
    title(doc, "Joint Journey", 22, after=2)
    title(doc, "Intended Use Statement and Safety Summary", 14, color=GREY, after=8)
    meta(doc, [
        ("Version:", "0.3 (draft)"),
        ("Date:", "18/07/2026"),

        ("Manufacturer:", "Elan Health Ltd (company number 17347255)"),
        ("Clinical Safety Officer:", "Mr Benjamin Zucker (registered CSO)"),

    ])

    # =====================================================================
    # PART 1 — INTENDED USE STATEMENT
    # =====================================================================
    title(doc, "Part 1: Intended Use Statement", 16)

    title(doc, "Primary Users (who directly interacts with the system)", 12)
    bullets(doc, [
        "Patients: adults on a waiting list for elective hip or knee replacement "
        "(primary user, self-directed).",
        "Carers and family members: may assist the user.",
        "Not a clinician-facing tool, although surgeons, GPs and physiotherapists "
        "may recommend it.",
    ])

    title(doc, "Technical Environment", 12)
    bullets(doc, [
        "Standalone web application (progressive web app).",
        "Runs on a mobile device and on desktop or laptop via a web browser.",
        "Not NHS Spine connected. Not integrated with an EPR. Not hosted on "
        "NHSmail infrastructure.",
    ])

    title(doc, "Purpose (what does it do?)", 12)
    bullets(doc, [
        "Primary function: structured digital prehabilitation, comprising guided "
        "strengthening and range-of-movement exercise, weight-management guidance, "
        "psychological and pain-preparation content, and pre-surgery education.",
        "Problem it solves: long UK waiting times mean patients often deteriorate "
        "(deconditioning, weight gain, anxiety) while they wait. Joint Journey uses "
        "that waiting period productively.",
        "Clinical outcomes supported: improved muscular strength and joint function, "
        "a healthier pre-operative weight and BMI, and better psychological readiness, "
        "all of which are associated with reduced length of stay and better recovery "
        "after joint replacement.",
    ])

    title(doc, "Environment (where is it used?)", 12)
    para(doc, "Primarily the patient's home, unsupervised and self-paced.")

    title(doc, "Intended Use Statement", 12)
    statement(doc,
        "Joint Journey is intended to provide structured prehabilitation support "
        "(exercise, weight management, psychological preparation and education) for "
        "use by adults awaiting elective hip or knee replacement in their own home, "
        "unsupervised, to improve their physical and psychological readiness for "
        "surgery and support better post-operative recovery.")

    title(doc, "The system DOES / ENABLES", 12)
    bullets(doc, [
        "Delivers evidence-based, self-directed exercise programmes banded by the "
        "validated Oxford Hip and Knee Score.",
        "Provides general weight-management and nutrition guidance with a safe "
        "minimum-calorie floor.",
        "Offers supportive mental-wellbeing and pain-preparation content with crisis "
        "signposting.",
        "Educates users on what to expect before, during and after surgery.",
        "Lets users track their own progress (weight, exercise adherence, scores).",
    ])

    title(doc, "The system is NOT intended to", 12)
    bullets(doc, [
        "Diagnose any condition or replace assessment by the clinical team.",
        "Direct, alter or replace prescribed treatment or the surgical pathway.",
        "Provide individualised medical, physiotherapy, dietetic or psychological "
        "therapy.",
        "Function as a medical device for monitoring, alerting or clinical "
        "decision-making.",
        "Be used in place of urgent care. Users are signposted to their GP, 111 or "
        "999 for red-flag symptoms.",
    ])

    # =====================================================================
    # PART 2 — SAFETY SUMMARY
    # =====================================================================
    title(doc, "Part 2: Safety Summary", 16)

    title(doc, "Technical File and Historical Documentation Review", 12)
    para(doc,
        "The clinical safety documentation set has been reviewed: the Intended "
        "Purpose and Claims policy, the DCB0129 Hazard Log (H01 to H10) and the "
        "Clinical Safety Case Report. Joint Journey is a low-risk health and "

        "wellbeing product providing general guidance. On its current scope it is "
        "not a medical device (no diagnosis, monitoring or treatment function). "
        "Data protection is handled via a DPIA and DTAC (cross-referenced for "
        "hazard H07). As a new product there is no legacy incident history, and a "
        "post-market safety-monitoring process is defined.")

    title(doc, "Hazard Summary", 12)
    para(doc,
        "Ten hazards have been identified (H01 to H10), covering exercise safety, "
        "clinical substitution, programme banding, nutrition, mental health, "
        "software errors, data security, accessibility, content currency, and "
        "third-party service availability. After mitigation, all residual risks are "
        "within the acceptable range (score 2 to 6 on the 5x5 matrix). The two "
        "highest-severity hazards are H02 (reliance on the app instead of seeking "
        "help) and H05 (mental-health distress); both are reduced by explicit "
        "red-flag and crisis signposting that is now implemented in the app. No "
        "residual risk is rated unacceptable. The overall risk profile is low, "
        "consistent with a general prehabilitation and wellbeing tool.")


    title(doc, "Evidence Base (how the system was reviewed)", 12)
    para(doc,
        "Assessment was based on first-hand review. This included direct use of the "
        "live web application and its content modules, review of the exercise, "
        "nutrition and mindset content against published guidance, a "
        "feature-by-feature hazard identification using structured HAZID and SWIFT "
        "techniques, and planned clinical advisory board sign-off of each domain "
        "(physiotherapy, dietetics and psychology). Conclusions are drawn from the "
        "product itself, not from documentation alone.")

    title(doc, "Overall Governance", 12)
    para(doc,
        "Governance for a founder-led early-stage product is developing but "
        "appropriately structured. A registered Clinical Safety Officer (Mr Benjamin "
        "Zucker) is in place. A documented clinical risk "

        "management approach (Hazard Log and Safety Case Report) is in place. A "
        "clinical advisory board is being engaged for content sign-off. "
        "Version-controlled release checks and an in-app feedback route for safety "
        "issues both exist. Software development lifecycle, incident management and "
        "support arrangements are proportionate to a low-risk product and should be "
        "formalised further as the product scales.")

    title(doc, "Deployment Readiness and Conditions", 12)
    para(doc,
        "The system is assessed as safe for a limited, free soft launch or pilot, on "
        "the condition that the open safety actions below are closed first. Status "
        "at this version:")
    bullets_status(doc, [
        ("Explicit \u201cwhen to seek help\u201d red-flag guidance (999 / 111 / GP / "
         "surgical team) reachable via a persistent header control on every in-app "
         "view, opening a panel (H02).", "IMPLEMENTED"),

        ("Crisis signposting, including Samaritans, surfaced alongside the "
         "wellbeing content (H05).", "IMPLEMENTED"),
        ("Pre-exercise safety check displayed at the top of the exercise "
         "programme (H01).", "IMPLEMENTED"),
        ("Formal pre-start exercise safety screening questionnaire (H01).", "OPEN"),
        ("CSO appointed and trained (Mr Benjamin Zucker).", "DONE"),
        ("Nutrition safeguards: minimum-calorie floor and condition cautions "
         "(H04).", "OPEN"),
        ("WCAG 2.1 AA accessibility pass completed (H08).", "OPEN"),
        ("DPIA cross-referenced for H07 (data controller, data residency "
         "confirmed UK).", "DONE"),
        ("Clinical advisory board review and sign-off of all hazard domains "
         "(H01\u2013H10).", "OPEN"),
        ("Formal CSO sign-off of Hazard Log (version bump to v1.0).", "OPEN"),

    ])
    para(doc,
        "Wider rollout or any NHS deployment should follow the deploying "
        "organisation's own DCB0160 assessment.")

    title(doc, "Intended Use, Scope and Clinical Responsibility", 12)
    para(doc,
        "Joint Journey provides general prehabilitation guidance and education only. "
        "It does not diagnose, does not individualise treatment, and does not "
        "replace the clinical team. Clinical responsibility for the patient remains "
        "at all times with their treating clinicians (surgeon, GP and "
        "physiotherapist). The app is a self-management adjunct during the wait for "
        "surgery. There are no EPR or Spine integrations, so there are no "
        "dual-system data-reconciliation risks. The main user-facing risk to "
        "communicate is that the app must not be used as a substitute for seeking "
        "medical help.")

    doc.add_paragraph("")
    sign = doc.add_paragraph()
    sign.add_run("Clinical Safety Officer sign-off:  ").bold = True
    sign.add_run("Name ____________________   Signature ____________________   Date __________")

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
