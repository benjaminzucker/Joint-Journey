#!/usr/bin/env python3
"""Generate the Advisory Board Pack as Word documents."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import shutil, os

OUT = os.path.dirname(os.path.abspath(__file__))

# -- Shared styling helpers ------------------------------------------

GREEN = RGBColor(0x47, 0x59, 0x53)
DARK = RGBColor(0x1F, 0x2D, 0x28)
GREY = RGBColor(0x55, 0x55, 0x55)

def set_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    for level in range(1, 4):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.color.rgb = GREEN

def add_logo_header(doc, title):
    """Add a simple branded header."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('⛰  Joint Journey')
    run.font.size = Pt(18)
    run.font.color.rgb = GREEN
    run.font.bold = True
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(title)
    run2.font.size = Pt(14)
    run2.font.color.rgb = DARK
    run2.font.bold = True
    
    # Thin line
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(12)
    run3 = p3.add_run('-' * 72)
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

def add_field(doc, label, default=''):
    p = doc.add_paragraph()
    run = p.add_run(f'{label}: ')
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(default if default else '______________________________')
    return p

def save(doc, name):
    path = os.path.join(OUT, name)
    doc.save(path)
    print(f'  ✓ {name}')


# ════════════════════════════════════════════════════════════════════
# 1. COVER LETTER
# ════════════════════════════════════════════════════════════════════
def create_cover_letter():
    doc = Document()
    set_style(doc)
    add_logo_header(doc, 'Clinical Advisory Board - Welcome Pack')
    
    add_field(doc, 'Date', '[                    ]')
    add_field(doc, 'Dear', '[Advisor name]')
    
    doc.add_paragraph('')
    doc.add_paragraph(
        'Thank you for agreeing to join the Joint Journey Clinical Advisory Board. '
        'Your expertise will be invaluable in ensuring our prehabilitation programme '
        'is safe, evidence-based, and genuinely helpful for patients awaiting hip or '
        'knee replacement surgery.'
    )
    
    doc.add_paragraph(
        'This pack contains all the documents you need to read and sign before our '
        'first meeting. Please take the time to review each one carefully.'
    )
    
    doc.add_heading("What's in this pack", level=2)

    
    items = [
        ('1. FAST Agreement', 
         'Your formal advisor agreement, covering the advisory relationship, '
         'confidentiality, intellectual property, and equity compensation. '
         'Please sign and return.'),
        ('2. Terms of Reference', 
         'Defines the Board\'s purpose, membership, responsibilities, and how '
         'meetings will work. For your information - we will formally adopt '
         'these at the first Board meeting.'),
        ('3. Conflict of Interest Policy', 
         'Our policy on identifying and managing conflicts of interest. '
         'Please read before completing the declaration form.'),
        ('4. Conflict of Interest Declaration Form', 
         'Please complete and return this form, declaring any relevant interests '
         '(or confirming "nil to declare"). This is a regulatory expectation '
         'and protects both you and the company.'),
        ('5. Content Review & Sign-off Log', 
         'The log we will use to record your formal sign-off of clinical content '
         'in your area of expertise. For reference - you will use this as we '
         'develop content together.'),
        ('6. Reviewer Sign-off Statement', 
         'The individual sign-off statement you will complete when reviewing '
         'specific clinical content. For reference.'),
    ]
    
    for title, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p2 = doc.add_paragraph(desc)
        p2.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_heading('What you need to do', level=2)
    
    actions = [
        'Read all documents in this pack',
        'Sign and return the FAST Agreement',
        'Complete and return the Conflict of Interest Declaration Form',
        'Ensure your professional indemnity insurance covers private advisory work '
        '(this is your responsibility as an HCPC/GMC/BABCP registrant)',
        'Let me know your availability for our first Board meeting',
    ]
    for i, action in enumerate(actions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. ')
        run.bold = True
        p.add_run(action)
    
    doc.add_paragraph('')
    doc.add_paragraph(
        'If you have any questions about any of the documents, please don\'t '
        'hesitate to get in touch. I\'m genuinely excited to work with you on this.'
    )
    
    doc.add_paragraph('')
    doc.add_paragraph('With thanks,')
    p = doc.add_paragraph()
    run = p.add_run('Benjamin Zucker')
    run.bold = True
    doc.add_paragraph('Founder, Joint Journey')
    doc.add_paragraph('hello@jointjourney.org')
    
    save(doc, '1 - Cover Letter.docx')


# ════════════════════════════════════════════════════════════════════
# 2. TERMS OF REFERENCE
# ════════════════════════════════════════════════════════════════════
def create_tor():
    doc = Document()
    set_style(doc)
    add_logo_header(doc, 'Clinical Advisory Board - Terms of Reference')
    
    p = doc.add_paragraph()
    run = p.add_run('Version: ')
    run.bold = True
    p.add_run('0.1 (draft)')
    run2 = p.add_run('    Date: ')
    run2.bold = True
    p.add_run('[        ]')
    run3 = p.add_run('    Owner: ')
    run3.bold = True
    p.add_run('Mr Benjamin Zucker, Founder')
    
    # Section 1
    doc.add_heading('1. Purpose', level=2)
    doc.add_paragraph(
        'The Clinical Advisory Board ("the Board") provides independent, multidisciplinary '
        'clinical input to Joint Journey, a digital prehabilitation programme for adults '
        'awaiting hip or knee replacement. The Board exists to:'
    )
    purposes = [
        'Assure clinical content - ensure the exercise, nutrition and mental-health/pain-preparation content is safe, evidence-based and appropriate.',
        'Support clinical risk management - contribute to hazard identification and mitigation (feeding the DCB0129 clinical safety case).',
        'Guide claims - advise on what can and cannot be claimed for the product.',
        'Lend credibility - provide named clinical oversight for users, NHS trusts and funders.',
    ]
    for purpose in purposes:
        doc.add_paragraph(purpose, style='List Bullet')
    
    doc.add_paragraph(
        'The Board is advisory: it informs decisions but does not direct the company. '
        'Final decisions rest with the company/founder, except that clinical content '
        'must not be published without the relevant advisor\'s documented sign-off.'
    )
    
    # Section 2
    doc.add_heading('2. Membership', level=2)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ['Seat', 'Discipline', 'Holder']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    
    rows_data = [
        ('Chair', 'Founder / surgical oversight (orthopaedics)', 'Mr Benjamin Zucker'),
        ('Member', 'MSK / orthopaedic Physiotherapist (HCPC)', '[                ]'),
        ('Member', 'Registered Dietitian (HCPC)', '[                ]'),
        ('Member', 'Clinical / Health Psychologist (HCPC) or CBT therapist (BABCP)', '[                ]'),
        ('Member (recommended)', 'Patient representative (PPIE) - lived experience of joint replacement', '[                ]'),
        ('Optional', 'GP / pre-operative nurse', '[                ]'),
    ]
    for i, (seat, disc, holder) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = seat
        table.rows[i].cells[1].text = disc
        table.rows[i].cells[2].text = holder
    
    doc.add_paragraph('')
    doc.add_paragraph('All clinical members must hold current professional registration (HCPC/GMC/BABCP).')
    doc.add_paragraph('Members serve a renewable term of 2 years.')
    doc.add_paragraph(
        'The Clinical Safety Officer (CSO) for the product is Mr Benjamin Zucker. '
        'The Board supports, but does not replace, the CSO.'
    )
    
    # Section 3
    doc.add_heading('3. Roles & Responsibilities', level=2)
    p = doc.add_paragraph()
    run = p.add_run('Members will:')
    run.bold = True
    
    will_do = [
        'Review the content for their pillar (initial and on material change) and record sign-off in the Content Review & Sign-off Log.',
        'Help identify hazards and mitigations in their domain.',
        'Declare conflicts of interest.',
        'Attend meetings (or send comments) and act within their professional scope.',
    ]
    for item in will_do:
        doc.add_paragraph(item, style='List Bullet')
    
    p = doc.add_paragraph()
    run = p.add_run('Members will NOT:')
    run.bold = True
    
    wont_do = [
        'Take responsibility for company decisions or operations.',
        'Provide individual clinical care to users through this role.',
    ]
    for item in wont_do:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 4
    doc.add_heading('4. Independence & Conflicts of Interest', level=2)
    doc.add_paragraph(
        'Members should be independent of any NHS organisation that may later commission or '
        'purchase Joint Journey, to avoid procurement conflicts. Where a member has such a '
        'connection, it must be declared and managed (see Conflict of Interest Policy).'
    )
    doc.add_paragraph(
        'Conflicts are declared on joining, recorded in the register, and revisited at every meeting.'
    )
    
    # Section 5
    doc.add_heading('5. Meetings', level=2)
    meetings = [
        ('Frequency:', 'Quarterly (≈4/year), plus async reviews as content is developed.'),
        ('Quorum:', 'The Chair plus at least 2 members; the relevant pillar\'s advisor must be present (or have commented) for any decision on that pillar.'),
        ('Minutes:', 'Taken for every meeting, recording attendance, decisions, sign-offs, actions and declared interests.'),
        ('Decisions:', 'By consensus; the Chair holds a casting view on non-clinical matters. Clinical content sign-off rests with the relevant registered advisor.'),
    ]
    for label, text in meetings:
        p = doc.add_paragraph()
        run = p.add_run(label + ' ')
        run.bold = True
        p.add_run(text)
    
    # Section 6
    doc.add_heading('6. Compensation', level=2)
    doc.add_paragraph('As set out in each member\'s FAST Agreement.')
    
    # Section 7
    doc.add_heading('7. Confidentiality & IP', level=2)
    doc.add_paragraph('Members keep company information confidential (see FAST Agreement).')
    doc.add_paragraph('Intellectual property in content/advice created for the company is assigned to the company (see FAST Agreement).')
    
    # Section 8
    doc.add_heading('8. Indemnity & Responsibility', level=2)
    doc.add_paragraph('Members provide advice in good faith within their professional competence.')
    doc.add_paragraph(
        'The company (via the CSO) holds responsibility for the product and its deployment. '
        'Members are responsible for ensuring their own professional indemnity insurance '
        'covers their advisory activities.'
    )
    
    # Section 9
    doc.add_heading('9. Review', level=2)
    doc.add_paragraph(
        'This ToR is reviewed annually and updated as the product and regulatory position evolve.'
    )
    
    doc.add_paragraph('')
    doc.add_paragraph('-' * 50)
    add_field(doc, 'Adopted by the Board on')
    add_field(doc, 'Chair signature')
    
    save(doc, '3 - Terms of Reference.docx')


# ════════════════════════════════════════════════════════════════════
# 3. CONFLICT OF INTEREST POLICY
# ════════════════════════════════════════════════════════════════════
def create_coi_policy():
    doc = Document()
    set_style(doc)
    add_logo_header(doc, 'Conflict of Interest Policy')
    
    p = doc.add_paragraph()
    run = p.add_run('Version: ')
    run.bold = True
    p.add_run('0.1 (draft)')
    run2 = p.add_run('    Date: ')
    run2.bold = True
    p.add_run('[        ]')
    
    doc.add_paragraph(
        'Applies to: Clinical Advisory Board members, the founder, and anyone advising '
        'or working on Joint Journey.'
    )
    
    doc.add_heading('1. Purpose', level=2)
    doc.add_paragraph(
        'To identify, declare and manage any interests that could (or could appear to) '
        'improperly influence the advice given to, or decisions made by, Joint Journey.'
    )
    
    doc.add_heading('2. What is a conflict of interest?', level=2)
    doc.add_paragraph('An interest that might affect, or be seen to affect, a person\'s objectivity. Types:')
    
    types = [
        ('Financial', 'payments, equity, shares, consultancy with competitors or suppliers.'),
        ('Non-financial / personal', 'friendships, family relationships, reputation.'),
        ('Loyalty / role', 'employment or office at an NHS organisation that may commission or buy the product; roles with competing products.'),
        ('Indirect', 'interests of a close family member or close associate.'),
    ]
    for label, desc in types:
        p = doc.add_paragraph()
        run = p.add_run(f'{label} - ')
        run.bold = True
        p.add_run(desc)
    
    doc.add_heading('3. Examples relevant to us', level=2)
    examples = [
        'An advisor employed by a trust that might pilot or purchase Joint Journey.',
        'An advisor with shares in, or paid by, a competing prehab/physio product.',
        'A personal friendship between the founder and an advisor (declare it - it does not bar the role, but should be on record).',
    ]
    for ex in examples:
        doc.add_paragraph(ex, style='List Bullet')
    
    doc.add_heading('4. Duties', level=2)
    doc.add_paragraph('Everyone covered by this policy must:')
    duties = [
        'Declare relevant interests on joining, using the Conflict of Interest Declaration Form.',
        'Update the declaration promptly when interests change.',
        'Disclose at the start of each meeting any interest relevant to the agenda.',
        'Step back from any decision where they have a material conflict (the minutes must record this).',
    ]
    for i, duty in enumerate(duties, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. ')
        run.bold = True
        p.add_run(duty)
    
    doc.add_heading('5. Managing a conflict', level=2)
    doc.add_paragraph('Depending on severity, the Board may:')
    actions = [
        'Record and note it (minor / declared only); or',
        'Exclude the person from the relevant discussion/decision; or',
        'Remove the relevant item from their remit; or',
        'In serious cases, end the advisory relationship.',
    ]
    for action in actions:
        doc.add_paragraph(action, style='List Bullet')
    
    doc.add_paragraph(
        'The Chair decides how each conflict is managed and records it in the meeting minutes.'
    )
    
    doc.add_heading('6. Review', level=2)
    doc.add_paragraph(
        'This policy is reviewed annually and updated as the company and its NHS relationships develop.'
    )
    
    doc.add_paragraph('')
    doc.add_paragraph('-' * 50)
    add_field(doc, 'Approved by')
    add_field(doc, 'Date')
    
    save(doc, '4 - Conflict of Interest Policy.docx')


# ════════════════════════════════════════════════════════════════════
# 4. CONFLICT OF INTEREST DECLARATION FORM
# ════════════════════════════════════════════════════════════════════
def create_coi_form():
    doc = Document()
    set_style(doc)
    add_logo_header(doc, 'Conflict of Interest Declaration Form')
    
    doc.add_paragraph(
        'Please complete this form and return it before your first Advisory Board meeting. '
        'If you have no interests to declare, please tick "Nil to declare" at the bottom. '
        'Please refer to the Conflict of Interest Policy for definitions and examples.'
    )
    
    doc.add_heading('Your details', level=2)
    add_field(doc, 'Name')
    add_field(doc, 'Role on Advisory Board')
    add_field(doc, 'Professional registration (e.g. HCPC PH12345)')
    add_field(doc, 'Employer')
    add_field(doc, 'Date')
    
    doc.add_heading('Declarations', level=2)
    doc.add_paragraph(
        'Please declare any interests that could (or could appear to) influence your '
        'advisory role with Joint Journey. If in doubt, declare it - it is always better '
        'to over-declare than under-declare.'
    )
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ['Type of interest', 'Details', 'How should this be managed?']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    
    types = [
        'Financial (e.g. paid work for competitor, shares)',
        'Non-financial / personal (e.g. friendship with founder)',
        'Loyalty / role (e.g. NHS employer who may buy the product)',
        'Indirect (e.g. family member\'s interests)',
    ]
    for i, t in enumerate(types):
        table.rows[i+1].cells[0].text = t
        table.rows[i+1].cells[1].text = ''
        table.rows[i+1].cells[2].text = ''
    
    doc.add_paragraph('')
    
    # Nil declaration
    p = doc.add_paragraph()
    run = p.add_run('☐  ')
    run.font.size = Pt(14)
    run2 = p.add_run('Nil to declare - ')
    run2.bold = True
    p.add_run(
        'I confirm that I have no interests to declare at this time. '
        'I understand my obligation to update this declaration if my circumstances change.'
    )
    
    doc.add_paragraph('')
    doc.add_paragraph('-' * 50)
    doc.add_paragraph('')
    add_field(doc, 'Signed')
    add_field(doc, 'Print name')
    add_field(doc, 'Date')
    
    save(doc, '5 - Conflict of Interest Declaration Form.docx')


# ════════════════════════════════════════════════════════════════════
# COPY EXISTING DOCUMENTS
# ════════════════════════════════════════════════════════════════════
def copy_existing():
    src_dir = os.path.join(os.path.dirname(OUT), 'advisory-board')
    copies = [
        ('FI FAST Agreement.docx', '2 - FAST Agreement.docx'),
        ('content-review-and-signoff-log.docx', '6 - Content Review and Sign-off Log.docx'),
        ('reviewer-signoff-statement.docx', '7 - Reviewer Sign-off Statement.docx'),
    ]
    for src_name, dst_name in copies:
        src = os.path.join(src_dir, src_name)
        dst = os.path.join(OUT, dst_name)
        if os.path.exists(src):
            shutil.copy2(src, dst)
            print(f'  ✓ {dst_name} (copied)')
        else:
            print(f'  ✗ {src_name} not found at {src}')


# ════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('Generating Advisory Board Pack...\n')
    create_cover_letter()
    create_tor()
    create_coi_policy()
    create_coi_form()
    copy_existing()
    print(f'\nDone! Files saved to:\n  {OUT}')
