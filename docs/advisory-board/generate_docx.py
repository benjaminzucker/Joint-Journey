#!/usr/bin/env python3
"""
Generate the two Word documents for the Joint Journey Clinical Advisory Board:
  1. content-review-and-signoff-log.docx
  2. reviewer-signoff-statement.docx

Usage:
    python3 generate_docx.py
Requires: python-docx  (pip3 install python-docx)
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))

# Product components reviewers may sign off (edit to match the live product)
COMPONENTS = [
    ("Exercise programme — Hip", "Physiotherapy"),
    ("Exercise programme — Knee", "Physiotherapy"),
    ("Range-of-movement guidance", "Physiotherapy"),
    ("Weight management / nutrition content", "Dietetics"),
    ("Mental health & pain-preparation content", "Psychology"),
    ("Getting ready for surgery content", "Multidisciplinary"),
    ("Oxford Hip/Knee Score implementation", "Multidisciplinary"),
    ("'Why It Works' / evidence content", "Multidisciplinary"),
]

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, size=16, color=NAVY, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.space_after = Pt(space_after)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY
    return p


# ---------------------------------------------------------------------------
# 1) Content review & sign-off log
# ---------------------------------------------------------------------------
def build_log():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    add_heading(doc, "Joint Journey — Clinical Content Review & Sign-off Log", 16)
    add_note(
        doc,
        "Evidence that clinical content has been reviewed and approved by an "
        "appropriately registered clinician. Supports the DCB0129 clinical safety "
        "case and any NHS DTAC submission. One row per component, per review.",
    )

    meta = doc.add_paragraph()
    meta.add_run("Product / version: ").bold = True
    meta.add_run("Joint Journey [version] \u2003")
    meta.add_run("Log owner (CSO): ").bold = True
    meta.add_run("[Mr Benjamin Zucker]")

    doc.add_paragraph("")

    headers = [
        "Component",
        "Pillar",
        "Version reviewed",
        "Reviewer (name + HCPC/GMC no.)",
        "Date",
        "Comments / changes requested",
        "Outcome (Approved / Approved w/ changes / Not approved)",
        "Re-review due",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=8, color=RGBColor(0xFF, 0xFF, 0xFF))
        # shade header
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F3A5F")
        tcPr.append(shd)

    for comp, pillar in COMPONENTS:
        row = table.add_row().cells
        set_cell_text(row[0], comp, size=8)
        set_cell_text(row[1], pillar, size=8)
        for j in range(2, len(headers)):
            set_cell_text(row[j], "", size=8)

    # column widths
    widths = [1.4, 0.9, 0.9, 1.4, 0.7, 1.6, 1.4, 0.8]
    for r in table.rows:
        for idx, w in enumerate(widths):
            r.cells[idx].width = Inches(w)

    doc.add_paragraph("")
    add_note(
        doc,
        "Guidance: 'Version reviewed' ties the sign-off to a specific content "
        "version. When content materially changes, add a NEW row (do not overwrite) "
        "and set a re-review date. File the matching signed statement "
        "(reviewer-signoff-statement.docx) and attribute any hazards in the hazard log.",
    )

    out = os.path.join(HERE, "content-review-and-signoff-log.docx")
    doc.save(out)
    print("Wrote", out)


# ---------------------------------------------------------------------------
# 2) Reviewer sign-off statement
# ---------------------------------------------------------------------------
def build_statement():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "Joint Journey — Clinical Reviewer Sign-off Statement", 16)
    add_note(
        doc,
        "One statement per reviewer, per content version reviewed. Email return is "
        "acceptable as a record (a wet signature is not required).",
    )

    doc.add_paragraph("")

    def field(label, value=""):
        p = doc.add_paragraph()
        p.add_run(label + " ").bold = True
        p.add_run(value if value else "_______________________________________")

    field("Reviewer name:")
    field("Profession & registration body:", "[e.g. Physiotherapist, HCPC]")
    field("Registration number:")
    field("Advisory role / seat:")
    field("Component(s) reviewed:")
    field("Content version reviewed:")
    field("Date of review:")

    doc.add_paragraph("")
    add_heading(doc, "Declaration", 12)

    decl = doc.add_paragraph()
    decl.add_run(
        "I confirm that I have reviewed the content listed above, that it falls "
        "within my professional competence, and that — subject to any comments "
        "noted below — I consider it clinically appropriate and reasonable for use "
        "in the Joint Journey prehabilitation programme for adults awaiting hip or "
        "knee replacement. I understand this content is educational and supportive "
        "and is not a substitute for individual clinical assessment."
    ).font.size = Pt(11)

    doc.add_paragraph("")
    o = doc.add_paragraph()
    o.add_run("Outcome (circle one): ").bold = True
    o.add_run("Approved   /   Approved with changes   /   Not approved")

    doc.add_paragraph("")
    c = doc.add_paragraph()
    c.add_run("Comments / required changes:").bold = True
    for _ in range(4):
        doc.add_paragraph("_______________________________________________________________")

    doc.add_paragraph("")
    doc.add_paragraph("")
    sig = doc.add_paragraph()
    sig.add_run("Signed: ").bold = True
    sig.add_run("____________________________   ")
    sig.add_run("Date: ").bold = True
    sig.add_run("__________________")

    doc.add_paragraph("")
    add_note(
        doc,
        "Conflicts of interest: please ensure any relevant interests are recorded in "
        "the Conflict of Interest Register before signing.",
    )

    out = os.path.join(HERE, "reviewer-signoff-statement.docx")
    doc.save(out)
    print("Wrote", out)


if __name__ == "__main__":
    build_log()
    build_statement()
    print("Done.")
