#!/usr/bin/env python3
"""
Generate: IP Assignment Agreement (Founder → Elan Health Ltd).

Produces a valid, Word-openable .docx.  Run:

    python3 generate_ip_assignment_docx.py

Requires: python-docx  (pip3 install python-docx)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "IP Assignment Agreement - Elan Health Ltd.docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
BLACK = RGBColor(0x00, 0x00, 0x00)


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def para(doc, text, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p


def numbered(doc, number, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{number}. ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def signature_block(doc, title, name_line):
    doc.add_paragraph("")
    p1 = doc.add_paragraph()
    p1.add_run(title).bold = True
    p1.paragraph_format.space_after = Pt(4)

    for label in ["Signed: ", "Name:  ", "Date:  "]:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.font.size = Pt(11)
        r.bold = True
        r2 = p.add_run("_" * 40 if label.startswith("Signed") else
                        (name_line if label.startswith("Name") else "_" * 30))
        r2.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("INTELLECTUAL PROPERTY ASSIGNMENT AGREEMENT")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    t.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Between Benjamin Zucker and Elan Health Ltd")
    r2.font.size = Pt(12)
    r2.font.color.rgb = NAVY
    sub.paragraph_format.space_after = Pt(16)

    # ---- Date & parties ----
    para(doc, "Date: ____________________", bold=True, after=12)

    heading(doc, "Parties", level=2)
    numbered(doc, 1,
        "Benjamin Zucker of 154 Beesmoor Road, Frampton Cotterell, Bristol, "
        "BS36 2JN (the \"Assignor\")")
    numbered(doc, 2,
        "Elan Health Ltd, a company registered in England and Wales under "
        "company number 17347255, whose registered office is at 154 Beesmoor "
        "Road, Frampton Cotterell, Bristol, BS36 2JN (the \"Company\")")

    # ---- Background ----
    heading(doc, "Background", level=2)
    para(doc,
        "A. The Assignor has developed a digital prehabilitation product known "
        "as \"Joint Journey\" (the \"Product\"), including its software source "
        "code, website, written content, exercise programmes, recipes, "
        "educational and wellbeing materials, visual design, brand identity, "
        "logo, and all associated documentation.")
    para(doc,
        "B. The Assignor is the sole founder and director of the Company and "
        "wishes to assign all intellectual property rights in the Product to "
        "the Company so that it is owned by and for the benefit of the Company.")

    # ---- Operative provisions ----
    heading(doc, "Agreed Terms", level=2)

    heading(doc, "1. Definitions", level=3)
    para(doc,
        "\"Intellectual Property Rights\" or \"IP\" means all copyright, database "
        "rights, design rights (registered and unregistered), trade marks and "
        "trade mark applications, domain name registrations, know-how, "
        "confidential information, and all other intellectual property rights "
        "of any kind, in each case whether registered or unregistered and "
        "including all applications and rights to apply for any of the foregoing, "
        "anywhere in the world.")
    para(doc,
        "\"Product\" means the digital health product known as \"Joint Journey\", "
        "including but not limited to:")
    bullet(doc, "All source code (HTML, CSS, JavaScript, configuration files, build scripts)")
    bullet(doc, "All written content (exercise descriptions, nutritional guidance, "
               "mindset modules, educational articles, recipes)")
    bullet(doc, "All visual design, user interface design, icons, illustrations, and layout")
    bullet(doc, "The brand name \"Joint Journey\", the logo, and all associated branding")
    bullet(doc, "The domain name jointjourney.org and any related domain registrations")
    bullet(doc, "All clinical safety documentation, hazard logs, and regulatory materials")
    bullet(doc, "All business plans, financial models, roadmaps, and strategic documents")
    bullet(doc, "Any other materials created by the Assignor in connection with the Product "
               "prior to the date of this agreement")

    heading(doc, "2. Assignment", level=3)
    para(doc,
        "2.1 In consideration of the sum of one pound (£1.00) paid by the "
        "Company to the Assignor (receipt of which the Assignor acknowledges), "
        "the Assignor hereby assigns to the Company, with full title guarantee, "
        "all Intellectual Property Rights in the Product, to hold absolutely "
        "for the full period of such rights and all renewals and extensions "
        "thereof.")
    para(doc,
        "2.2 The assignment in clause 2.1 includes all Intellectual Property "
        "Rights created by the Assignor prior to and up to the date of this "
        "agreement.")
    para(doc,
        "2.3 From the date of this agreement, any Intellectual Property Rights "
        "created by the Assignor in connection with the Product shall vest in "
        "the Company automatically by virtue of his role as director and "
        "employee/officer of the Company.")

    heading(doc, "3. Warranties", level=3)
    para(doc, "The Assignor warrants that:")
    bullet(doc, "He is the sole owner of the Intellectual Property Rights in the Product")
    bullet(doc, "The IP is free from any encumbrance, lien, licence, or third-party claim")
    bullet(doc, "He has not previously assigned or licensed the IP to any other party")
    bullet(doc, "To the best of his knowledge, the Product does not infringe the "
               "intellectual property rights of any third party")
    bullet(doc, "He has full power and authority to enter into this agreement")

    heading(doc, "4. Further Assurance", level=3)
    para(doc,
        "The Assignor shall, at the Company's cost, execute all documents and "
        "do all things reasonably necessary to give full effect to this "
        "agreement, including assisting with any trade mark applications, "
        "domain transfers, or registrations.")

    heading(doc, "5. Moral Rights", level=3)
    para(doc,
        "To the extent permitted by law, the Assignor waives any and all moral "
        "rights he may have in the Product under the Copyright, Designs and "
        "Patents Act 1988 (or any equivalent legislation).")

    heading(doc, "6. Governing Law", level=3)
    para(doc,
        "This agreement is governed by and shall be construed in accordance "
        "with the laws of England and Wales. The courts of England and Wales "
        "shall have exclusive jurisdiction.")

    # ---- Execution ----
    heading(doc, "Execution", level=2)
    para(doc,
        "This agreement has been entered into on the date stated above.",
        after=12)

    signature_block(doc,
        "Signed by BENJAMIN ZUCKER (Assignor):",
        "Benjamin Zucker")

    signature_block(doc,
        "Signed by a director of ELAN HEALTH LTD (Company):",
        "Benjamin Zucker, Director")

    # ---- Save ----
    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
